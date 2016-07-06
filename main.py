import re

from docx.text.run import Run
from lxml.etree import Element, tostring

from docx import Document
from mailmerge import MailMerge

NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'ct': 'http://schemas.openxmlformats.org/package/2006/content-types',
}


def preprocess(tree):
    r = re.compile(r' MERGEFIELD +"?([^ ]+?)"? +(|\\\* MERGEFORMAT )', re.I)
    # Remove attribute that soft-links to other namespaces; other namespaces
    # are not used, so would cause word to throw an error.
    ignorable_key = '{%(mc)s}Ignorable' % NAMESPACES
    if ignorable_key in tree.getroot().attrib:
        del tree.getroot().attrib[ignorable_key]

    to_delete = []

    for parent in tree.iterfind('.//{%(w)s}fldSimple/..' % NAMESPACES):
        for idx, child in enumerate(parent):
            if child.tag != '{%(w)s}fldSimple' % NAMESPACES:
                continue
            instr = child.attrib['{%(w)s}instr' % NAMESPACES]

            m = r.match(instr)
            if not m:
                raise ValueError('Could not determine name of merge '
                                 'field in value "%s"' % instr)
            parent[idx] = Element('MergeField', name=m.group(1))

    for parent in tree.iterfind('.//{%(w)s}instrText/../..' % NAMESPACES):
        children = list(parent)
        fields = zip(
            [children.index(e) for e in
             parent.findall('{%(w)s}r/{%(w)s}fldChar[@{%(w)s}fldCharType="begin"]/..' % NAMESPACES)],
            [children.index(e) for e in
             parent.findall('{%(w)s}r/{%(w)s}fldChar[@{%(w)s}fldCharType="end"]/..' % NAMESPACES)],
            [e.text for e in
             parent.findall('{%(w)s}r/{%(w)s}instrText' % NAMESPACES)])
        for idx_begin, idx_end, instr in fields:
            m = r.match(instr)
            if m is None:
                continue
            parent[idx_begin] = Element('MergeField', name=m.group(1))
            to_delete += [(parent, parent[i + 1])
                          for i in range(idx_begin, idx_end)]

    for parent, child in to_delete:
        parent.remove(child)


doc = Document('examples/complex_img.docx')

preprocess(doc._element.getroottree())
for par in doc.paragraphs:
    fields = par._element.findall('MergeField[@name="image"]')
    for field in fields:
        run = Run(par._p.add_r(), par)
        run.add_picture('../pigeony.jpg')
        par._element.replace(field, run._element)

doc.save('examples/generated_img.docx')

doc = MailMerge('../complex_table.docx')

context = {
    'foo': 'Standards are confusing.',
    'bar': '56',
    'baz': 'hop',
    'col1': [
        {'col1': 'hop', 'col2': '1', 'col3': '4', 'col4': 'FOUR'},
        {'col1': 'la', 'col2': '2', 'col3': '5', 'col4': 'STILL FOUR'},
        {'col1': 'kee', 'col2': '3', 'col3': '6', 'col4': 'FOUR'},
    ]
}

print(doc.get_merge_fields())

doc.merge(**context)
print(tostring(list(doc.parts.values())[0], pretty_print=True, encoding='utf-8').decode('utf-8'))
doc.write('examples/generated_table.docx')

